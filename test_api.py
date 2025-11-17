"""
Test script for the DOCX to PDF API endpoint
NOTE: Requires 'requests' library: pip install requests
"""
import os
import tempfile

try:
    import requests
except ImportError:
    print("ERROR: 'requests' library not installed.")
    print("Please run: pip install requests")
    exit(1)

def create_test_docx():
    """Create a simple test DOCX file using python-docx if available, or use a minimal valid DOCX"""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test document for API testing.')
        doc.add_paragraph('It contains some sample text to verify conversion works.')
        
        temp_path = os.path.join(tempfile.gettempdir(), 'test_document.docx')
        doc.save(temp_path)
        return temp_path
    except ImportError:
        print("NOTE: python-docx not installed. Please provide a test .docx file manually.")
        print("For this test, please place a DOCX file at: C:\\temp\\test.docx")
        test_path = r"C:\temp\test.docx"
        if os.path.exists(test_path):
            return test_path
        else:
            raise FileNotFoundError(f"Please create a test DOCX file at: {test_path}")

def test_api_endpoint(url='http://127.0.0.1:5005/api/convert'):
    """Test the API endpoint"""
    print(f"Testing API endpoint: {url}")
    print("-" * 50)
    
    # Create test DOCX
    print("Creating test DOCX file...")
    docx_path = create_test_docx()
    print(f"✓ Created: {docx_path}")
    
    # Test the API
    print(f"\nSending POST request to {url}...")
    with open(docx_path, 'rb') as f:
        files = {'file': f}
        response = requests.post(url, files=files)
    
    print(f"Status Code: {response.status_code}")
    print(f"Content-Type: {response.headers.get('Content-Type')}")
    
    if response.status_code == 200:
        # Save the PDF
        output_path = os.path.join(tempfile.gettempdir(), 'test_output.pdf')
        with open(output_path, 'wb') as pdf_file:
            pdf_file.write(response.content)
        
        print(f"✓ SUCCESS: PDF created ({len(response.content)} bytes)")
        print(f"✓ Saved to: {output_path}")
        return True
    else:
        print(f"✗ FAILED: {response.text}")
        return False

def test_error_handling(url='http://127.0.0.1:5005/api/convert'):
    """Test error handling"""
    print(f"\n\nTesting error handling...")
    print("-" * 50)
    
    # Test 1: No file
    print("\nTest 1: No file uploaded")
    response = requests.post(url)
    print(f"Status: {response.status_code} - {response.json()}")
    
    # Test 2: Wrong file type
    print("\nTest 2: Wrong file type (txt instead of docx)")
    temp_txt = os.path.join(tempfile.gettempdir(), 'test.txt')
    with open(temp_txt, 'w') as f:
        f.write('test')
    
    with open(temp_txt, 'rb') as f:
        files = {'file': ('test.txt', f)}
        response = requests.post(url, files=files)
    print(f"Status: {response.status_code} - {response.json()}")
    
    os.remove(temp_txt)

if __name__ == '__main__':
    print("DOCX to PDF API Test")
    print("=" * 50)
    
    try:
        # Test main functionality
        success = test_api_endpoint()
        
        # Test error handling
        test_error_handling()
        
        print("\n" + "=" * 50)
        if success:
            print("✓ ALL TESTS PASSED")
        else:
            print("✗ SOME TESTS FAILED")
    
    except Exception as e:
        print(f"\n✗ ERROR: {e}")
        import traceback
        traceback.print_exc()
